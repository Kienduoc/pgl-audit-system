from docx import Document
import sys

doc_path = r'C:\Users\Admin\.gemini\antigravity\playground\triple-disk\pgl-audit-system\Document\2. PGL-TT21-BM01b-Phieu_Dang_Ky_Chung_Nhan_San_Pham_PT5.docx'

try:
    doc = Document(doc_path)
    
    print("=" * 80)
    print("DOCUMENT CONTENT")
    print("=" * 80)
    
    # Print all paragraphs
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"\n[Para {i}]: {para.text}")
    
    # Print all tables
    print("\n" + "=" * 80)
    print("TABLES")
    print("=" * 80)
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n--- Table {table_idx + 1} ---")
        for row_idx, row in enumerate(table.rows):
            cells_text = [cell.text.strip() for cell in row.cells]
            print(f"Row {row_idx}: {' | '.join(cells_text)}")
            
except Exception as e:
    print(f"Error: {e}")
    sys.exit(1)
