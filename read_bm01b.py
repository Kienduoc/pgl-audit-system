import sys
import io
from docx import Document

# Force UTF-8 for stdout
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc_path = r'C:\Users\Admin\.gemini\antigravity\playground\triple-disk\pgl-audit-system\Document\2. PGL-TT21-BM01b-Phieu_Dang_Ky_Chung_Nhan_San_Pham_PT5.docx'

try:
    doc = Document(doc_path)
    print("--- DOCUMENT START ---")
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text.strip())
    
    print("\n--- TABLES ---")
    for t in doc.tables:
        for r in t.rows:
            row_text = []
            for c in r.cells:
                row_text.append(c.text.strip())
            print(" | ".join(row_text))
    print("--- DOCUMENT END ---")

except Exception as e:
    print(f"Error: {e}")
