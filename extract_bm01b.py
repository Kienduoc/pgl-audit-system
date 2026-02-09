import sys
import io
from docx import Document

# Set recursion limit just in case
sys.setrecursionlimit(2000)

doc_path = r'C:\Users\Admin\.gemini\antigravity\playground\triple-disk\pgl-audit-system\Document\2. PGL-TT21-BM01b-Phieu_Dang_Ky_Chung_Nhan_San_Pham_PT5.docx'
output_path = r'C:\Users\Admin\.gemini\antigravity\playground\triple-disk\pgl-audit-system\bm01b_full_content.txt'

try:
    doc = Document(doc_path)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("=== DOCUMENT START ===\n\n")
        
        # Paragraphs
        f.write("--- PARAGRAPHS ---\n")
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                f.write(f"P{i}: {text}\n")
        
        # Tables
        f.write("\n--- TABLES ---\n")
        for i, table in enumerate(doc.tables):
            f.write(f"\n[Table {i+1}]\n")
            for j, row in enumerate(table.rows):
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                f.write(f"Row {j}: {' | '.join(cells)}\n")
                
        f.write("\n=== DOCUMENT END ===\n")
        
    print(f"Successfully wrote content to {output_path}")

except Exception as e:
    print(f"Error: {e}")
